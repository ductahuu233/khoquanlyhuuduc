import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from app.services.word_maker import generate_word

def inspect_word():
    test_data = {
        "request_id": 2,
        "requester_name": "Đại úy Phạm Văn B (Phòng Kỹ Thuật)",
        "destination": "Đội Cảnh Sát Giao Thông Số 1 - Công An Tỉnh",
        "reason": "Trang bị phục vụ công tác tuần tra kiểm soát giao thông khẩn cấp",
        "export_date": "26/07/2026",
        "items": [
            {"item_code": "VT-CA01", "name": "Bộ Đàm Motorola GP328 Chuyên Dụng", "unit": "Bộ", "quantity": 5},
            {"item_code": "VT-CA02", "name": "Đèn Pin Siêu Sáng Đuốc Tuần Tra", "unit": "Cái", "quantity": 10}
        ]
    }

    path = generate_word(test_data)
    doc = Document(path)

    print("=== PARAGRAPHS IN GENERATED WORD FILE ===")
    for i, p in enumerate(doc.paragraphs):
        print(f"P{i}: '{p.text}' | Runs: {[(r.text, r.font.color.rgb, r.bold) for r in p.runs]}")

    print("\n=== TABLES IN GENERATED WORD FILE ===")
    print(f"Total tables: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx} (rows: {len(table.rows)}, cols: {len(table.columns)}):")
        for r_idx, row in enumerate(table.rows):
            cell_texts = [cell.text.replace('\n', ' ') for cell in row.cells]
            print(f"  Row {r_idx}: {cell_texts}")

if __name__ == "__main__":
    inspect_word()
