import os
from datetime import datetime
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.database import get_db
from app.models import Item, Asset, AssetHistory, log_audit
from app.schemas import AssetReportDamagedPayload, AssetTransferPayload
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

router = APIRouter(prefix="/api/assets", tags=["Asset Status & Lifecycle"])

EXPORTS_DIR = os.path.join(os.getcwd(), "exports")

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 8)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

@router.get("")
def get_assets(status_filter: Optional[str] = None, location_filter: Optional[str] = None, db: Session = Depends(get_db)):
    query = db.query(Asset)
    if status_filter:
        query = query.filter(Asset.status == status_filter)
    if location_filter:
        query = query.filter(Asset.location == location_filter)
    
    assets = query.order_by(Asset.id.desc()).all()
    return [
        {
            "id": a.id,
            "asset_code": a.asset_code,
            "item_code": a.item.item_code if a.item else "",
            "name": a.item.name if a.item else "Tài sản",
            "unit": a.item.unit if a.item else "Cái",
            "serial_number": a.serial_number or "N/A",
            "mac_address": a.mac_address or "N/A",
            "status": a.status,
            "assigned_to": a.assigned_to or "Kho Kỹ Thuật",
            "location": a.location,
            "created_at": a.created_at.isoformat()
        } for a in assets
    ]

@router.get("/scan/{code}")
def scan_code(code: str, db: Session = Depends(get_db)):
    """Quét mã QR (Hỗ trợ cả Mã Tem TS-xxxx và Mã Vật Tư VT-xxxx)"""
    asset = db.query(Asset).filter(Asset.asset_code == code.upper()).first()
    if asset:
        return {
            "type": "fixed_asset",
            "id": asset.id,
            "code": asset.asset_code,
            "name": asset.item.name if asset.item else "",
            "item_code": asset.item.item_code if asset.item else "",
            "serial_number": asset.serial_number,
            "status": asset.status,
            "assigned_to": asset.assigned_to,
            "location": asset.location
        }
    
    item = db.query(Item).filter(Item.item_code == code.upper()).first()
    if item:
        return {
            "type": "consumable",
            "id": item.id,
            "code": item.item_code,
            "name": item.name,
            "unit": item.unit,
            "current_stock": item.current_stock,
            "category": item.category,
            "location": item.location
        }
    
    raise HTTPException(
        status_code=status.HTTP_404_NOT_FOUND,
        detail=f"Không tìm thấy dữ liệu vật tư/tài sản có mã QR: {code}"
    )

@router.post("/report-damaged")
def report_asset_damaged(payload: AssetReportDamagedPayload, db: Session = Depends(get_db)):
    """Cán bộ quét mã QR trên di động và báo hỏng hóc thiết bị"""
    asset = db.query(Asset).filter(Asset.asset_code == payload.asset_code.upper()).first()
    if not asset:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Không tìm thấy tài sản mã {payload.asset_code}"
        )
    
    old_status = asset.status
    asset.status = "damaged"
    
    history = AssetHistory(
        item_id=asset.item_id,
        asset_id=asset.id,
        action_type="damaged",
        performer=payload.reporter_name or "Cán bộ sử dụng",
        details=f"Báo hỏng hóc thiết bị từ trạng thái [{old_status}]. Lý do: {payload.reason}"
    )
    db.add(history)
    db.commit()

    log_audit(db, "user", "BÁO HỎNG THIẾT BỊ", asset.asset_code, f"Lý do: {payload.reason}")
    return {
        "success": True,
        "message": f"Đã ghi nhận báo hỏng cho thiết bị {asset.asset_code} ({asset.item.name if asset.item else ''}). Chờ thu hồi về Kho Phế Phẩm."
    }

@router.post("/recover")
def recover_damaged_asset(asset_code: str, performer: str = "Thủ kho", db: Session = Depends(get_db)):
    """Thủ kho nhận lại đồ hỏng, xác nhận đưa vào Kho Phế Phẩm"""
    asset = db.query(Asset).filter(Asset.asset_code == asset_code.upper()).first()
    if not asset:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài sản")

    asset.status = "damaged"
    asset.location = "Kho Phế Phẩm"

    history = AssetHistory(
        item_id=asset.item_id,
        asset_id=asset.id,
        action_type="recover",
        performer=performer,
        details="Thủ kho xác nhận thu hồi thiết bị hỏng đưa vào Kho Phế Phẩm chờ thanh lý"
    )
    db.add(history)
    db.commit()

    log_audit(db, "storekeeper", "THU HỒI PHẾ PHẨM", asset.asset_code, "Xác nhận đưa vào Kho Phế Phẩm")
    return {"success": True, "message": f"Đã thu hồi tài sản {asset.asset_code} vào Kho Phế Phẩm"}

@router.post("/transfer")
def transfer_asset(payload: AssetTransferPayload, db: Session = Depends(get_db)):
    """Điều chuyển tài sản giữa các phòng/cán bộ"""
    asset = db.query(Asset).filter(Asset.asset_code == payload.asset_code.upper()).first()
    if not asset:
        raise HTTPException(status_code=404, detail="Không tìm thấy tài sản")

    old_location = asset.location
    old_assigned = asset.assigned_to

    asset.location = payload.to_location
    asset.assigned_to = payload.assigned_to
    asset.status = "in_use"

    history = AssetHistory(
        item_id=asset.item_id,
        asset_id=asset.id,
        action_type="transfer",
        performer=payload.performer or "Thủ kho",
        details=f"Điều chuyển từ [{old_location} - {old_assigned}] sang [{payload.to_location} - {payload.assigned_to}]"
    )
    db.add(history)
    db.commit()

    return {"success": True, "message": f"Đã điều chuyển tài sản {asset.asset_code} sang {payload.assigned_to}"}

@router.get("/disposal-word")
def generate_disposal_word(db: Session = Depends(get_db)):
    """
    Sinh Tờ Trình / Biên Bản Thanh Lý Tài Sản A4 (.docx) cho Kho Phế Phẩm chuẩn Nghị định 30.
    """
    damaged_assets = db.query(Asset).filter(Asset.status.in_(["damaged", "disposed"])).all()
    
    if not os.path.exists(EXPORTS_DIR):
        os.makedirs(EXPORTS_DIR, exist_ok=True)
    
    file_path = os.path.join(EXPORTS_DIR, "bien_ban_thanh_ly_tai_san.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.79)
        section.bottom_margin = Inches(0.79)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(0.79)

    # Header Table
    h_table = doc.add_table(rows=1, cols=2)
    h_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_left, c_right = h_table.cell(0, 0), h_table.cell(0, 1)

    p_l = c_left.paragraphs[0]
    p_l.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l1 = p_l.add_run("BỘ TƯ LỆNH CẢNH SÁT CƠ ĐỘNG\n")
    r_l1.bold = True
    r_l1.font.size = Pt(10)
    r_l2 = p_l.add_run("ĐOÀN NGHI LỄ CÔNG AN NHÂN DÂN\n")
    r_l2.bold = True
    r_l2.font.size = Pt(10.5)
    p_l.add_run("Số: ....../TTr-TL").font.size = Pt(11)

    p_r = c_right.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_r1.bold = True
    r_r1.font.size = Pt(11)
    r_r2 = p_r.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r_r2.bold = True
    r_r2.font.size = Pt(11)
    p_r.add_run("......, ngày ..... tháng ..... năm 2026").italic = True

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(14)
    run_t = p_title.add_run("TỜ TRÌNH XIN THANH LÝ TÀI SẢN PHẾ PHẨM")
    run_t.bold = True
    run_t.font.size = Pt(15)

    p_kg = doc.add_paragraph()
    p_kg.paragraph_format.space_before = Pt(12)
    r_kg = p_kg.add_run("Kính gửi: Ban Chỉ huy Đoàn Nghi lễ CAND / Bộ Tư lệnh Cảnh sát Cơ động")
    r_kg.bold = True

    p_body = doc.add_paragraph()
    p_body.paragraph_format.space_before = Pt(8)
    p_body.add_run("Kính trình Lãnh đạo phê duyệt danh mục tài sản hỏng hóc không còn khả năng sử dụng đưa vào thanh lý phế phẩm dưới đây:\n")

    # Table of damaged assets
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["STT", "Mã Tem QR", "Tên Thiết Bị / Tài Sản", "Số Serial", "Tình Trạng Hỏng"]
    widths = [Inches(0.6), Inches(1.2), Inches(2.6), Inches(1.2), Inches(1.2)]
    b_black = {"val": "single", "sz": 8, "color": "000000"}

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = widths[i]
        cell.text = h
        set_cell_border(cell, top=b_black, bottom=b_black, left=b_black, right=b_black)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(11)

    for idx, asset in enumerate(damaged_assets, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = asset.asset_code
        row_cells[2].text = asset.item.name if asset.item else ""
        row_cells[3].text = asset.serial_number or "N/A"
        row_cells[4].text = "Hỏng nặng / Chờ thanh lý"

        for i, c in enumerate(row_cells):
            c.width = widths[i]
            set_cell_border(c, top=b_black, bottom=b_black, left=b_black, right=b_black)
            p = c.paragraphs[0]
            if i in [0, 1, 3, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Khung chữ ký chuẩn NĐ 30
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(14)
    p_space.paragraph_format.space_after = Pt(10)

    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'none')
                tcBorders.append(node)
            tcPr.append(tcBorders)

    s_left, s_right = sig_table.cell(0, 0), sig_table.cell(0, 1)
    s_left.width, s_right.width = Inches(3.2), Inches(3.4)

    p_t = s_left.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t1 = p_t.add_run("TRƯỞNG ĐOÀN\n")
    r_t1.bold = True
    r_t1.font.size = Pt(12)
    r_t2 = p_t.add_run("(Ký, đóng dấu và ghi rõ họ tên)\n\n\n\n\n")
    r_t2.italic = True
    r_t2.font.size = Pt(10)

    p_r = s_right.paragraphs[0]
    p_r.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_r1 = p_r.add_run("NGƯỜI TRÌNH\n")
    r_r1.bold = True
    r_r1.font.size = Pt(12)
    r_r2 = p_r.add_run("(Ký và ghi rõ họ tên)\n\n\n\n\n")
    r_r2.italic = True
    r_r2.font.size = Pt(10)

    doc.save(file_path)
    return FileResponse(
        path=file_path,
        filename="to_trinh_thanh_ly_tai_san.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
