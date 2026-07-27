import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

OUTPUT_DIR = os.path.join(os.getcwd(), "exports")

def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
    """Thiết lập lề trong cho ô bảng Word."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """
    Gán viền đen nét đậm cho ô bảng.
    """
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 8))) # 1pt solid line
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def generate_word(export_data: dict) -> str:
    """
    Sinh Tờ trình xuất kho vật tư chuẩn thể thức hành chính Nghị định 30/2020/NĐ-CP (.docx).
    """
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR, exist_ok=True)

    request_id = export_data.get("request_id", "000")
    file_path = os.path.join(OUTPUT_DIR, f"to_trinh_xuat_kho_{request_id}.docx")

    doc = Document()

    # Cấu hình lề trang A4 chuẩn Nghị định 30 (Trên 2cm, Dưới 2cm, Trái 3cm, Phải 2cm)
    for section in doc.sections:
        section.top_margin = Inches(0.79)     # 2 cm
        section.bottom_margin = Inches(0.79)  # 2 cm
        section.left_margin = Inches(1.18)    # 3 cm
        section.right_margin = Inches(0.79)   # 2 cm

    # 1. KHUNG TIÊU NGỮ VÀ ĐƠN VỊ BAN HÀNH (Bảng 2 cột không viền chuẩn NĐ 30)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False

    # Bỏ viền bảng header
    for row in header_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'none')
                tcBorders.append(node)
            tcPr.append(tcBorders)

    c_left = header_table.cell(0, 0)
    c_right = header_table.cell(0, 1)
    c_left.width = Inches(2.8)
    c_right.width = Inches(3.6)

    # Cột trái: Tên cơ quan & Số hiệu
    p_left = c_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_cq1 = p_left.add_run("BỘ CÔNG AN\n")
    r_cq1.font.name = "Times New Roman"
    r_cq1.font.size = Pt(11)
    
    r_cq2 = p_left.add_run("CỤC KỸ THUẬT VẬT TƯ\n")
    r_cq2.bold = True
    r_cq2.font.name = "Times New Roman"
    r_cq2.font.size = Pt(11)

    r_so = p_left.add_run(f"Số: {request_id}/TTr-PXK")
    r_so.font.name = "Times New Roman"
    r_so.font.size = Pt(11)

    # Cột phải: Quốc hiệu & Tiêu ngữ & Ngày tháng
    p_right = c_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_qh = p_right.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r_qh.bold = True
    r_qh.font.name = "Times New Roman"
    r_qh.font.size = Pt(11)

    r_tn = p_right.add_run("Độc lập - Tự do - Hạnh phúc\n")
    r_tn.bold = True
    r_tn.font.name = "Times New Roman"
    r_tn.font.size = Pt(11)

    # Trích xuất ngày tháng chuẩn (bỏ giờ phút giây)
    raw_date_str = str(export_data.get("export_date", ""))
    try:
        if " " in raw_date_str:
            date_part = raw_date_str.split(" ")[0]
            dt = datetime.strptime(date_part, "%d/%m/%Y")
        else:
            dt = datetime.strptime(raw_date_str, "%d/%m/%Y")
    except Exception:
        dt = datetime.now()

    date_formatted = f"Hà Nội, ngày {dt.day:02d} tháng {dt.month:02d} năm {dt.year}"
    r_date = p_right.add_run(date_formatted)
    r_date.italic = True
    r_date.font.name = "Times New Roman"
    r_date.font.size = Pt(11)

    # Khoảng cách sau Header
    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)
    p_space.paragraph_format.space_after = Pt(6)

    # 2. TIÊU ĐỀ TỜ TRÌNH VÀ TRÍCH YẾU (MÀU ĐEN CHUẨN, IN HOA, CĂN GIỮA)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("TỜ TRÌNH")
    run_title.bold = True
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(15)
    run_title.font.color.rgb = RGBColor(0, 0, 0)  # Chuẩn màu đen NĐ 30

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    run_sub = p_sub.add_run("V/v xin phê duyệt xuất kho vật tư phục vụ công tác chuyên môn")
    run_sub.bold = True
    run_sub.font.name = "Times New Roman"
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(0, 0, 0)

    # 3. KÍNH GỬI
    requester = export_data.get("requester_name", "Đại úy Nguyễn Văn A")
    destination = export_data.get("destination", "Đơn vị tiếp nhận")
    reason = export_data.get("reason", "Phục vụ công tác chuyên môn")
    items = export_data.get("items", [])
    total_qty = sum(item.get("quantity", 0) for item in items)

    p_kg = doc.add_paragraph()
    p_kg.paragraph_format.space_after = Pt(12)
    p_kg.paragraph_format.line_spacing = 1.25
    r_kg_label = p_kg.add_run("Kính gửi: ")
    r_kg_label.bold = True
    r_kg_label.font.name = "Times New Roman"
    r_kg_label.font.size = Pt(13)
    p_kg.add_run("Ban Giám đốc / Thủ trưởng Cục Kỹ thuật Vật tư").font.name = "Times New Roman"

    # 4. NỘI DUNG TỜ TRÌNH (Trình bày văn phong hành chính, KHÔNG dùng dấu chấm bi bullet)
    p_body = doc.add_paragraph()
    p_body.paragraph_format.line_spacing = 1.25
    p_body.paragraph_format.space_after = Pt(8)

    # Mục 1: Căn cứ đề xuất
    r_sec1 = p_body.add_run("1. Căn cứ và lý do đề xuất:\n")
    r_sec1.bold = True
    r_sec1.font.name = "Times New Roman"
    r_sec1.font.size = Pt(13)
    p_body.add_run(f"- Căn cứ vào yêu cầu trang bị vật tư phục vụ công tác của: {destination}.\n").font.name = "Times New Roman"
    p_body.add_run(f"- Mục đích xuất kho: {reason}.\n\n").font.name = "Times New Roman"

    # Mục 2: Người đề xuất
    r_sec2 = p_body.add_run("2. Thông tin cán bộ đề xuất:\n")
    r_sec2.bold = True
    r_sec2.font.name = "Times New Roman"
    r_sec2.font.size = Pt(13)
    p_body.add_run(f"- Họ và tên cán bộ đề xuất: {requester}\n\n").font.name = "Times New Roman"

    # Mục 3: Chi tiết vật tư đề xuất
    r_sec3 = p_body.add_run("3. Nội dung đề xuất xuất kho:\n")
    r_sec3.bold = True
    r_sec3.font.name = "Times New Roman"
    r_sec3.font.size = Pt(13)
    p_body.add_run("Kính trình Lãnh đạo xem xét, phê duyệt xuất kho số lượng vật tư cụ thể theo danh sách dưới đây:").font.name = "Times New Roman"

    # 5. BẢNG DANH SÁCH VẬT TƯ (BẮT BUỘC KẺ VIỀN KÍN TOÀN BỘ CÁC Ô CHUẨN NĐ 30 - Table Grid)
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Ép viền bảng ở cấp độ Table XML Properties (<w:tblBorders>)
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')  # 1pt solid black line
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

    hdr_cells = table.rows[0].cells
    headers = ["STT", "Mã Vật Tư", "Tên Vật Tư / Thiết Bị Kho", "Đơn Vị Tính", "Số Lượng"]
    widths = [Inches(0.6), Inches(1.2), Inches(2.8), Inches(0.9), Inches(0.9)]

    b_black = {"val": "single", "sz": 8, "color": "000000"}

    for i, title in enumerate(headers):
        hdr_cells[i].width = widths[i]
        hdr_cells[i].text = title
        set_cell_border(hdr_cells[i], top=b_black, bottom=b_black, left=b_black, right=b_black)
        set_cell_margins(hdr_cells[i], top=120, bottom=120)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.bold = True
            r.font.size = Pt(11)

    # Tô màu nền tiêu đề bảng
    shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
    table.rows[0]._element.get_or_add_trPr().append(shading_elm)

    # Điền dòng vật tư
    for idx, item in enumerate(items, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = str(item.get("item_code", ""))
        row_cells[2].text = str(item.get("name", ""))
        row_cells[3].text = str(item.get("unit", ""))
        row_cells[4].text = str(item.get("quantity", 0))

        alignments = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        for i, cell in enumerate(row_cells):
            cell.width = widths[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(cell, top=b_black, bottom=b_black, left=b_black, right=b_black)
            set_cell_margins(cell, top=100, bottom=100)
            p = cell.paragraphs[0]
            p.alignment = alignments[i]
            for r in p.runs:
                r.font.name = "Times New Roman"
                r.font.size = Pt(11)

    # Dòng Tổng cộng
    total_row = table.add_row().cells
    total_row[0].text = ""
    total_row[1].text = ""
    total_row[2].text = "TỔNG CỘNG"
    total_row[3].text = ""
    total_row[4].text = str(total_qty)

    for cell in total_row:
        set_cell_border(cell, top=b_black, bottom=b_black, left=b_black, right=b_black)
        set_cell_margins(cell, top=100, bottom=100)

    p_tot_label = total_row[2].paragraphs[0]
    p_tot_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p_tot_label.runs:
        r.font.name = "Times New Roman"
        r.font.bold = True

    p_tot_val = total_row[4].paragraphs[0]
    p_tot_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_tot_val.runs:
        r.font.name = "Times New Roman"
        r.font.bold = True

    # 6. ĐOẠN KẾT
    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_before = Pt(14)
    p_end.paragraph_format.space_after = Pt(20)
    r_end = p_end.add_run("Kính trình Lãnh đạo đơn vị xem xét, phê duyệt./.")
    r_end.italic = True
    r_end.font.name = "Times New Roman"
    r_end.font.size = Pt(12)

    # 7. KHUNG CHỮ KÝ (Nơi nhận bên trái, Người trình bên phải)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Bỏ viền bảng chữ ký
    for row in sig_table.rows:
        for cell in row.cells:
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for b in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                node = OxmlElement(f'w:{b}')
                node.set(qn('w:val'), 'none')
                tcBorders.append(node)
            tcPr.append(tcBorders)

    s_left = sig_table.cell(0, 0)
    s_right = sig_table.cell(0, 1)
    s_left.width = Inches(3.0)
    s_right.width = Inches(3.4)

    # Bên trái: Nơi nhận
    p_nn = s_left.paragraphs[0]
    p_nn.paragraph_format.line_spacing = 1.15
    r_nn_title = p_nn.add_run("NƠI NHẬN:\n")
    r_nn_title.bold = True
    r_nn_title.font.name = "Times New Roman"
    r_nn_title.font.size = Pt(10)
    
    r_nn_body = p_nn.add_run("- Như trên;\n- Lưu: VT, Kho.")
    r_nn_body.font.name = "Times New Roman"
    r_nn_body.font.size = Pt(10)

    # Bên phải: Người trình
    p_trinh = s_right.paragraphs[0]
    p_trinh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_trinh1 = p_trinh.add_run("NGƯỜI TRÌNH\n")
    r_trinh1.bold = True
    r_trinh1.font.name = "Times New Roman"
    r_trinh1.font.size = Pt(12)

    r_trinh2 = p_trinh.add_run("(Ký và ghi rõ họ tên)\n\n\n\n")
    r_trinh2.italic = True
    r_trinh2.font.name = "Times New Roman"
    r_trinh2.font.size = Pt(10)

    r_trinh3 = p_trinh.add_run(f"{requester}")
    r_trinh3.bold = True
    r_trinh3.font.name = "Times New Roman"
    r_trinh3.font.size = Pt(12)

    # Khung phê duyệt của Lãnh đạo ở cuối
    p_app = doc.add_paragraph()
    p_app.paragraph_format.space_before = Pt(30)
    r_app = p_app.add_run("Ý KIẾN PHÊ DUYỆT CỦA LÃNH ĐẠO:\n")
    r_app.bold = True
    r_app.font.name = "Times New Roman"
    r_app.font.size = Pt(11)
    p_app.add_run("......................................................................................................................................................................\n").font.name = "Times New Roman"
    p_app.add_run("......................................................................................................................................................................").font.name = "Times New Roman"

    # Save file Word
    try:
        doc.save(file_path)
        return file_path
    except Exception as e:
        raise RuntimeError(f"Lỗi khi tạo file Word: {e}")
